#!/usr/bin/env python3
"""
Academic Report Generator — Markdown → Strictly Formatted .docx

Converts docs/academic-report.md into a monochromatic Chinese academic Word
document meeting national journal / thesis formatting standards.

Usage:
    python scripts/generate_docs.py
Output:
    Campus_Secondhand_Blockchain_Academic_Report.docx
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ═══════════════════════════════════════════════════════════════════════
# Constants
# ═══════════════════════════════════════════════════════════════════════

PROJECT_ROOT = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_ROOT / "docs" / "academic-report.md"
OUTPUT_PATH = PROJECT_ROOT / "Campus_Secondhand_Blockchain_Academic_Report.docx"

# Font names
FONT_BODY_CN = "宋体"          # SimSun
FONT_HEADING_CN = "黑体"       # SimHei
FONT_BODY_EN = "Times New Roman"
FONT_CODE = "Consolas"

# Sizes
SIZE_TITLE = Pt(22)            # 二号
SIZE_H1 = Pt(16)               # 三号
SIZE_H2 = Pt(14)               # 四号
SIZE_H3 = Pt(12)               # 小四
SIZE_BODY = Pt(12)             # 小四
SIZE_ABSTRACT = Pt(9)          # 小五
SIZE_TABLE = Pt(10.5)          # 五号
SIZE_CODE = Pt(9)              # 9pt
SIZE_FOOTNOTE = Pt(9)          # 小五

# Margins
MARGIN_TOP_BOTTOM = Cm(2.54)   # 1 inch
MARGIN_LEFT_RIGHT = Cm(3.18)   # 1.25 inch

LINE_SPACING_1_5 = 1.5

CODE_BG_COLOR = "F5F5F5"
BLACK = RGBColor(0, 0, 0)


# ═══════════════════════════════════════════════════════════════════════
# Low-level XML helpers
# ═══════════════════════════════════════════════════════════════════════

def _set_run_fonts(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_BODY):
    """Set both East-Asian and Latin fonts on a run."""
    run.font.size = size
    run.font.name = en_font
    run.font.color.rgb = BLACK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), en_font)


def _set_paragraph_spacing(para, line_spacing=LINE_SPACING_1_5,
                           before=0, after=0, first_line_indent=None):
    """Configure paragraph spacing and indentation."""
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def _add_left_border_to_paragraph(para, color="999999"):
    """Add a 1px gray left border to a paragraph (for code blocks)."""
    pPr = para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '4')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _add_shading_to_paragraph(para, color=CODE_BG_COLOR):
    """Add background shading to a paragraph."""
    pPr = para._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


def _make_three_line_table(table):
    """
    Transform a python-docx table into a classic three-line academic table:
    - Top of header: 1.5pt
    - Bottom of header: 0.75pt
    - Bottom of table: 1.5pt
    - Remove all vertical and internal horizontal borders
    - White background, no shading
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # Remove existing borders
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)

    tblBorders = OxmlElement('w:tblBorders')

    # Top border (1.5pt) — applied to all cells but we'll handle per-row
    # Instead, set all borders to none, then add specific ones
    for edge_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge = OxmlElement(f'w:{edge_name}')
        edge.set(qn('w:val'), 'none')
        edge.set(qn('w:sz'), '0')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), '000000')
        tblBorders.append(edge)

    tblPr.append(tblBorders)

    # Now set per-cell borders for three-line effect
    num_rows = len(table.rows)
    num_cols = len(table.columns)

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Remove existing cell borders
            for existing in tcPr.findall(qn('w:tcBorders')):
                tcPr.remove(existing)

            tcBorders = OxmlElement('w:tcBorders')

            if r_idx == 0:
                # Header row: top thick (1.5pt = 12 eighths), bottom thin (0.75pt = 6)
                top_edge = OxmlElement('w:top')
                top_edge.set(qn('w:val'), 'single')
                top_edge.set(qn('w:sz'), '12')  # 1.5pt
                top_edge.set(qn('w:space'), '0')
                top_edge.set(qn('w:color'), '000000')
                tcBorders.append(top_edge)

                bottom_edge = OxmlElement('w:bottom')
                bottom_edge.set(qn('w:val'), 'single')
                bottom_edge.set(qn('w:sz'), '6')  # 0.75pt
                bottom_edge.set(qn('w:space'), '0')
                bottom_edge.set(qn('w:color'), '000000')
                tcBorders.append(bottom_edge)

            elif r_idx == num_rows - 1:
                # Last row: bottom thick (1.5pt)
                bottom_edge = OxmlElement('w:bottom')
                bottom_edge.set(qn('w:val'), 'single')
                bottom_edge.set(qn('w:sz'), '12')
                bottom_edge.set(qn('w:space'), '0')
                bottom_edge.set(qn('w:color'), '000000')
                tcBorders.append(bottom_edge)

            # No left, right, or internal horizontal borders
            tcPr.append(tcBorders)

            # Remove cell shading
            tcPr_shd = tcPr.find(qn('w:shd'))
            if tcPr_shd is not None:
                tcPr.remove(tcPr_shd)


# ═══════════════════════════════════════════════════════════════════════
# Document setup
# ═══════════════════════════════════════════════════════════════════════

def setup_document():
    """Create a document with proper page layout and base style."""
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = MARGIN_TOP_BOTTOM
        section.bottom_margin = MARGIN_TOP_BOTTOM
        section.left_margin = MARGIN_LEFT_RIGHT
        section.right_margin = MARGIN_LEFT_RIGHT

    # --- Default style: body text ---
    style = doc.styles['Normal']
    style.font.name = FONT_BODY_EN
    style.font.size = SIZE_BODY
    style.font.color.rgb = BLACK
    style.paragraph_format.line_spacing = LINE_SPACING_1_5
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    # Set East-Asian font at style level
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_BODY_EN)
    rFonts.set(qn('w:hAnsi'), FONT_BODY_EN)
    rFonts.set(qn('w:eastAsia'), FONT_BODY_CN)

    return doc


# ═══════════════════════════════════════════════════════════════════════
# Block-level parsers / renderers
# ═══════════════════════════════════════════════════════════════════════

def add_title(doc, text):
    """Add the report title: 22pt, SimHei, centered, 12pt after."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=12)
    run = para.add_run(text)
    _set_run_fonts(run, cn_font=FONT_HEADING_CN, en_font=FONT_BODY_EN, size=SIZE_TITLE)
    run.bold = True


def add_abstract_heading(doc, text):
    """Abstract heading: 小五 宋体 bold."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0)
    run = para.add_run(text)
    _set_run_fonts(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_ABSTRACT)
    run.bold = True


def add_abstract_body(doc, text):
    """Abstract paragraph: 小五 宋体, justified, first-line indent 2 chars."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0,
                           first_line_indent=SIZE_ABSTRACT * 2)
    _add_rich_text(para, text, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_ABSTRACT)


def add_keywords(doc, text):
    """Keywords line: 小五 宋体, '关键词' bold."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0)
    # Split into label and content
    if "：" in text:
        label, content = text.split("：", 1)
        run_label = para.add_run(label + "：")
        _set_run_fonts(run_label, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_ABSTRACT)
        run_label.bold = True
        run_content = para.add_run(content)
        _set_run_fonts(run_content, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_ABSTRACT)
    else:
        run = para.add_run(text)
        _set_run_fonts(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_ABSTRACT)
        run.bold = True


def add_h1(doc, text):
    """一级标题: 三号(16pt), 黑体 bold, left aligned, 6pt before/after."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(para, line_spacing=1.5, before=6, after=6)
    run = para.add_run(text)
    _set_run_fonts(run, cn_font=FONT_HEADING_CN, en_font=FONT_BODY_EN, size=SIZE_H1)
    run.bold = True


def add_h2(doc, text):
    """二级标题: 四号(14pt), 黑体 bold, left aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(para, line_spacing=1.5, before=3, after=3)
    run = para.add_run(text)
    _set_run_fonts(run, cn_font=FONT_HEADING_CN, en_font=FONT_BODY_EN, size=SIZE_H2)
    run.bold = True


def add_h3(doc, text):
    """三级标题: 小四(12pt), 宋体 bold, left aligned."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(para, line_spacing=1.5, before=2, after=2)
    run = para.add_run(text)
    _set_run_fonts(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_H3)
    run.bold = True


def add_body_paragraph(doc, text):
    """Body paragraph: 小四(12pt), justified, first-line indent 2 chars."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0,
                           first_line_indent=SIZE_BODY * 2)
    _add_rich_text(para, text, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_BODY)


def add_code_block(doc, lines, language=None):
    """Code block: Consolas 9pt, gray background, left border, no indent, left margin 0.5in."""
    for line in lines:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(0)
        pf.left_indent = Inches(0.5)
        _add_left_border_to_paragraph(para)
        _add_shading_to_paragraph(para)
        run = para.add_run(line if line else " ")
        _set_run_fonts(run, cn_font=FONT_CODE, en_font=FONT_CODE, size=SIZE_CODE)
        run.font.name = FONT_CODE


def add_blockquote(doc, text):
    """Blockquote: 小四, italic-like treatment, left indent 0.5in."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = para.paragraph_format
    pf.line_spacing = LINE_SPACING_1_5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.left_indent = Inches(0.5)
    _add_rich_text(para, text, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_BODY)


def add_bullet_item(doc, text, indent_level=0):
    """Bullet list item with first-line indent 2 chars."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0,
                           first_line_indent=SIZE_BODY * 2)
    _add_rich_text(para, text, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_BODY)


def add_ordered_item(doc, text, indent_level=0):
    """Numbered list item."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_spacing(para, line_spacing=1.5, before=0, after=0,
                           first_line_indent=SIZE_BODY * 2)
    _add_rich_text(para, text, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_BODY)


def add_table_from_md(doc, header_row, data_rows):
    """Create a three-line academic table from parsed markdown table data."""
    num_cols = len(header_row)
    num_rows = 1 + len(data_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.autofit = True

    # Header row
    for c_idx, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[c_idx]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.line_spacing = 1.0
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.first_line_indent = Pt(0)
        run = para.add_run(cell_text.strip())
        _set_run_fonts(run, cn_font=FONT_HEADING_CN, en_font=FONT_BODY_EN, size=SIZE_TABLE)
        run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = para.paragraph_format
            pf.line_spacing = 1.0
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.first_line_indent = Pt(0)
            run = para.add_run(cell_text.strip())
            _set_run_fonts(run, cn_font=FONT_BODY_CN, en_font=FONT_BODY_EN, size=SIZE_TABLE)

    _make_three_line_table(table)

    # Add a small spacer paragraph after the table
    spacer = doc.add_paragraph()
    _set_paragraph_spacing(spacer, line_spacing=1.0, before=2, after=2)
    spacer.paragraph_format.first_line_indent = Pt(0)

    return table


def add_empty_line(doc):
    """Add a small spacer line."""
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, line_spacing=1.0, before=2, after=2)
    para.paragraph_format.first_line_indent = Pt(0)


# ═══════════════════════════════════════════════════════════════════════
# Rich text (inline formatting) parser
# ═══════════════════════════════════════════════════════════════════════

def _add_rich_text(para, text, cn_font, en_font, size):
    """
    Parse inline markdown formatting in text and add runs to paragraph.

    Handles:
      - **bold**
      - `inline code`
      - LaTeX math: $...$ and $$...$$
      - plain text (with Chinese/English mixed)
    """
    # Patterns in order of priority
    # 1. Inline code: `...`
    # 2. Bold: **...**
    # 3. LaTeX display math (on its own — handled at block level)
    # 4. Regular text

    # Tokenize the text
    tokens = _tokenize_inline(text)
    for token_type, token_text in tokens:
        if token_type == 'bold':
            run = para.add_run(token_text)
            _set_run_fonts(run, cn_font=cn_font, en_font=en_font, size=size)
            run.bold = True
        elif token_type == 'code':
            run = para.add_run(token_text)
            _set_run_fonts(run, cn_font=FONT_CODE, en_font=FONT_CODE, size=Pt(9))
        elif token_type == 'math':
            # Render LaTeX as italic math expression
            run = para.add_run(token_text)
            _set_run_fonts(run, cn_font=cn_font, en_font=en_font, size=size)
            run.italic = True
        else:
            run = para.add_run(token_text)
            _set_run_fonts(run, cn_font=cn_font, en_font=en_font, size=size)


def _tokenize_inline(text):
    """
    Split text into (type, content) tuples for inline formatting.

    Types: 'text', 'bold', 'code', 'math'
    """
    tokens = []
    i = 0
    while i < len(text):
        # Inline code: `...`
        if text[i] == '`' and (i == 0 or text[i-1] != '\\'):
            end = text.find('`', i + 1)
            if end != -1:
                tokens.append(('code', text[i+1:end]))
                i = end + 1
                continue

        # Display math: $$...$$
        if text[i:i+2] == '$$' and (i == 0 or text[i-1] != '\\'):
            end = text.find('$$', i + 2)
            if end != -1:
                tokens.append(('math', text[i+2:end]))
                i = end + 2
                continue

        # Inline math: $...$
        if text[i] == '$' and (i == 0 or text[i-1] != '\\') and (i+1 < len(text) and text[i+1] != '$'):
            end = text.find('$', i + 1)
            if end != -1:
                tokens.append(('math', text[i+1:end]))
                i = end + 1
                continue

        # Bold: **...**
        if text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                tokens.append(('bold', text[i+2:end]))
                i = end + 2
                continue

        # Plain text — accumulate until next special char
        j = i
        while j < len(text):
            if text[j] == '`' and (j == 0 or text[j-1] != '\\'):
                break
            if text[j:j+2] == '$$' and (j == 0 or text[j-1] != '\\'):
                break
            if text[j] == '$' and (j == 0 or text[j-1] != '\\') and (j+1 < len(text) and text[j+1] != '$'):
                break
            if text[j:j+2] == '**':
                break
            j += 1

        tokens.append(('text', text[i:j]))
        i = j

    return tokens


# ═══════════════════════════════════════════════════════════════════════
# Markdown Parser
# ═══════════════════════════════════════════════════════════════════════

def parse_and_render(doc, md_text):
    """
    Parse markdown text and render into python-docx Document.

    State-machine based parser. States:
      - NORMAL, CODE_BLOCK, TABLE
    """
    lines = md_text.split('\n')
    i = 0
    n = len(lines)

    # Track whether we've processed the title
    title_processed = False
    # Track whether we're in the abstract zone
    in_abstract = False
    abstract_lines = []
    keywords_line = None

    while i < n:
        line = lines[i]

        # ── Horizontal rule (skip) ──
        if line.strip() == '---':
            i += 1
            # After abstract section, flush abstract
            if in_abstract:
                _flush_abstract(doc, abstract_lines, keywords_line)
                in_abstract = False
                abstract_lines = []
                keywords_line = None
            continue

        # ── Code block ──
        if line.strip().startswith('```'):
            lang = line.strip()[3:].strip() or None
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code_block(doc, code_lines, lang)
            continue

        # ── Table ──
        if '|' in line and line.strip().startswith('|'):
            table_lines = []
            while i < n and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse table: first line = header, second = separator, rest = data
            if len(table_lines) >= 2:
                header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # Skip separator line (|---|...)
                data = []
                for tl in table_lines[2:]:
                    data.append([c.strip() for c in tl.split('|')[1:-1]])
                add_table_from_md(doc, header, data)
            continue

        # ── Headings ──
        if line.startswith('# '):
            title_processed = True
            add_title(doc, line[2:].strip())
            i += 1
            continue

        if line.startswith('## '):
            heading_text = line[3:].strip()
            if heading_text == '摘要':
                in_abstract = True
                abstract_lines = []
                keywords_line = None
                i += 1
                continue
            else:
                add_h1(doc, heading_text)
                i += 1
                continue

        if line.startswith('### '):
            add_h2(doc, line[4:].strip())
            i += 1
            continue

        if line.startswith('#### '):
            add_h3(doc, line[5:].strip())
            i += 1
            continue

        # ── Blockquote ──
        if line.startswith('> '):
            add_blockquote(doc, line[2:].strip())
            i += 1
            continue

        # ── Unordered list ──
        if re.match(r'^- (\*\*.*|.*)', line.strip()):
            add_bullet_item(doc, re.sub(r'^- ', '', line.strip()))
            i += 1
            continue

        # ── Ordered list ──
        if re.match(r'^\d+\. ', line.strip()):
            add_ordered_item(doc, re.sub(r'^\d+\. ', '', line.strip()))
            i += 1
            continue

        # ── Empty line ──
        if line.strip() == '':
            i += 1
            continue

        # ── Regular paragraph ──
        # Check if this is in the abstract section
        if in_abstract:
            # Check if it's a keywords line
            if line.strip().startswith('**关键词**'):
                keywords_line = line.strip()
                i += 1
                continue
            else:
                abstract_lines.append(line.strip())
                i += 1
                continue

        # Accumulate multi-line paragraphs (lines separated by single newline)
        para_lines = [line.strip()]
        i += 1
        while i < n and lines[i].strip() != '' \
                and not lines[i].startswith('#') \
                and not lines[i].startswith('```') \
                and not lines[i].startswith('> ') \
                and not lines[i].startswith('|') \
                and not lines[i].startswith('- ') \
                and not re.match(r'^\d+\. ', lines[i].strip()) \
                and not lines[i].strip() == '---':
            para_lines.append(lines[i].strip())
            i += 1

        full_para = ' '.join(para_lines)
        add_body_paragraph(doc, full_para)

    # Flush abstract if at end of document
    if in_abstract:
        _flush_abstract(doc, abstract_lines, keywords_line)


def _flush_abstract(doc, abstract_lines, keywords_line):
    """Write accumulated abstract content."""
    if not abstract_lines:
        return
    add_abstract_heading(doc, '摘要')
    full_abstract = ' '.join(abstract_lines)
    add_abstract_body(doc, full_abstract)
    if keywords_line:
        # Extract text from **keywords** format
        kw_text = keywords_line.replace('**', '')
        add_keywords(doc, kw_text)


# ═══════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════

def main():
    if not MD_PATH.exists():
        print(f"Error: {MD_PATH} not found")
        return 1

    md_text = MD_PATH.read_text(encoding='utf-8')
    doc = setup_document()
    parse_and_render(doc, md_text)

    # Save
    doc.save(str(OUTPUT_PATH))
    file_size = OUTPUT_PATH.stat().st_size
    print(f"Generated: {OUTPUT_PATH}")
    print(f"File size: {file_size:,} bytes ({file_size/1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
